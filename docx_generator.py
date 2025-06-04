from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT # Per l'allineamento verticale nelle celle
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re
import os
from typing import List

from core.config import settings

# Colore per "BN Surveys Srls" nel footer (campiona dal tuo PDF per precisione)
FOOTER_BN_SURVEYS_COLOR = RGBColor(0x00, 0x5A, 0x9E) # Esempio Blu (da affinare)

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Imposta i margini interni di una cella (in dxa, 1 inch = 1440 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_name, m_val in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        mar_el = OxmlElement(f'w:{m_name}')
        mar_el.set(qn('w:w'), str(m_val))
        mar_el.set(qn('w:type'), 'dxa')
        tcMar.append(mar_el)
    tcPr.append(tcMar)

def remove_table_borders(table):
    """Rende invisibili tutti i bordi di una tabella."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border_el = OxmlElement(f'w:{border_name}')
                border_el.set(qn('w:val'), 'nil') # 'nil' o 'none' per nessun bordo
                # border_el.set(qn('w:sz'), '0') # Opzionale: dimensione a 0
                # border_el.set(qn('w:space'), '0') # Opzionale: spazio a 0
                # border_el.set(qn('w:color'), 'auto') # Opzionale: colore auto
                tcBorders.append(border_el)
            tcPr.append(tcBorders)

def add_logo_to_header(header, logo_path, width_inch=1.7):
    """Aggiunge un logo all'header, allineato a sinistra."""
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.clear() # Rimuove contenuto preesistente dal paragrafo di default dell'header
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Riduci al minimo lo spazio prima/dopo il paragrafo del logo nell'header
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    # fmt.line_spacing = 1.0 # Opzionale se l'immagine ha spazio intrinseco

    try:
        run = paragraph.add_run()
        # Aggiungiamo l'immagine. La posizione esatta (alto a sinistra estremo)
        # dipende anche dai margini della pagina e dell'header.
        # python-docx non offre controllo fine del posizionamento "fluttuante" dell'immagine.
        # L'allineamento del paragrafo a sinistra è il massimo che possiamo fare standard.
        run.add_picture(logo_path, width=Inches(width_inch))
    except FileNotFoundError:
        print(f"ERRORE: File logo non trovato in {logo_path}. Il logo non sarà aggiunto.")
    except Exception as e:
        print(f"ERRORE durante l'aggiunta del logo: {e}. Il logo non sarà aggiunto.")


def create_styled_docx(plain_text_report_content: str) -> io.BytesIO:
    if not isinstance(plain_text_report_content, str):
        raise TypeError(
            f"Expected plain_text_report_content to be a string, "
            f"but got {type(plain_text_report_content)}. "
            f"Please check the caller in app.py to ensure 'report_content' is a string instance."
        )
    document: Document = Document()

    # --- Impostazioni di Stile di Default per l'Intero Documento ---
    style = document.styles['Normal']
    font = style.font
    font.name = settings.DOCX_FONT_NAME
    font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = settings.DOCX_LINE_SPACING if hasattr(settings, 'DOCX_LINE_SPACING') else 1.5
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(settings.DOCX_SPACE_AFTER_PARAGRAPH if hasattr(settings, 'DOCX_SPACE_AFTER_PARAGRAPH') else 6)

    # --- Header con Logo su Ogni Pagina ---
    # Il path va risolto correttamente rispetto all'esecuzione dello script.
    # Se lo script è nella root del progetto e assets è una subdir:
    logo_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'assets', 'logos', 'bn-surveys-logo.png')
    if not os.path.exists(logo_path): # Fallback se il path relativo sopra non funziona (es. in contesti diversi)
        logo_path = "assets/logos/bn-surveys-logo.png"

    header = document.sections[0].header
    add_logo_to_header(header, logo_path, width_inch=1.7) # Larghezza logo (da PDF sembra circa 1.7-1.8 inches)

    # --- Elaborazione del Contenuto Testuale Generato dall'LLM ---
    lines: List[str] = plain_text_report_content.split('\n')

    recipient_block_active = True
    date_place_block_next = False
    subject_line_pattern = re.compile(r"^\s*Oggetto\s*:\s*(.*)", re.IGNORECASE)
    recipient_line_pattern = re.compile(r"^\s*(Spett\.le|Via|Viale|Piazza|Corso|[0-9]{5}\s+[A-Za-zÀ-ú\s]+\s*\([A-Z]{2}\))", re.IGNORECASE)
    date_place_line_pattern = re.compile(r"^\s*[A-Za-zÀ-ú\s]+,\s*\d{1,2}\s+[a-zà-ú]+\s+\d{4}\s*$", re.IGNORECASE)
    reference_line_pattern = re.compile(r"^\s*(Vs\. Rif\.|Polizza|Ns\. Rif\.)\s*:\s*(.*)") # Per i riferimenti tipo Vs. Rif.
    section_title_pattern = re.compile(r"^\s*([0-9]+)\s*–\s*([A-Z\sÀ-Ù'&]+)\s*$") # Es. "1 – DATI GENERALI"

    recipient_lines_count = 0
    MAX_RECIPIENT_LINES = 4 # Numero di righe per il blocco Spett.le
    
    is_first_content_paragraph_after_initial_blocks = True # Per gestire spazio prima del primo paragrafo narrativo
    initial_right_aligned_lines_count = 0 # Counter for the first three lines

    for line_num, line in enumerate(lines):
        stripped_line = line.strip()
        original_line = line

        # Force the first three lines to be right-aligned
        if initial_right_aligned_lines_count <= 3:
            p = document.add_paragraph(stripped_line if stripped_line else " ") # Add a space if line is empty
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fmt = p.paragraph_format
            # Minimal spacing, similar to recipient block styling
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 1.0
            if initial_right_aligned_lines_count == 0 and line_num == 0: # Only for the very first line of the document
                fmt.space_before = Pt(10) # Add some space after the header/logo
            
            initial_right_aligned_lines_count += 1
            continue

        # 1. Gestione Blocco Indirizzo "Spett.le" (allineato a destra)
        # This block will now only be considered for lines after the first three.
        # Ensure recipient_block_active is still checked.
        if recipient_block_active:
            is_current_line_recipient_type = recipient_line_pattern.match(stripped_line)
            if is_current_line_recipient_type and recipient_lines_count < MAX_RECIPIENT_LINES:
                p = document.add_paragraph(stripped_line)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                fmt = p.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                fmt.line_spacing = 1.0
                recipient_lines_count += 1
                if recipient_lines_count == 1: # Prima riga di QUESTO blocco destinatario
                    fmt.space_before = Pt(10) 
                continue # Riga processata come destinatario
            else: # La riga corrente NON è di tipo destinatario OPPURE si è raggiunto il max di righe per il blocco
                if recipient_lines_count > 0: # Se ERAVAMO in un blocco destinatario
                    recipient_block_active = False # Ora il blocco è ufficialmente terminato
                    date_place_block_next = True   # La prossima riga potrebbe essere data/luogo
                # Se recipient_lines_count == 0, significa che non abbiamo ancora trovato l'inizio
                # del blocco destinatario. recipient_block_active rimane True in questo caso.
                # La riga corrente (che non è di tipo destinatario) passerà ai controlli successivi.
        
        # Se recipient_block_active è False (perché terminato in una precedente iterazione), 
        # la logica sopra 'if recipient_block_active:' viene saltata, 
        # e le righe successive non verranno considerate per il blocco destinatario.

        # 2. Gestione Blocco Data e Luogo (allineato a sinistra)
        if date_place_block_next:
            is_date_place_line = date_place_line_pattern.match(stripped_line)
            if is_date_place_line:
                p = document.add_paragraph(stripped_line)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fmt = p.paragraph_format
                fmt.space_before = Pt(10) # Spazio dopo il blocco Spett.le
                fmt.space_after = Pt(6) # Spazio prima dei riferimenti
                date_place_block_next = False
                is_first_content_paragraph_after_initial_blocks = True # Reimposta per il prossimo blocco
                continue
            else:
                date_place_block_next = False # Non era una data, passa al contenuto normale

        # 3. Gestione Riferimenti (Vs. Rif., Polizza, Ns. Rif.)
        ref_match = reference_line_pattern.match(stripped_line)
        if ref_match:
            p = document.add_paragraph(stripped_line) # LLM dovrebbe aver allineato con spazi
            fmt = p.paragraph_format
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 1.0
            # Potresti voler un font leggermente più piccolo per questi, come nel PDF
            for run in p.runs:
                run.font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL) # Es. 10pt se normale è 11pt
                run.bold = True # Make reference lines bold
            if is_first_content_paragraph_after_initial_blocks:
                fmt.space_before = Pt(6) # Spazio prima del blocco riferimenti
                is_first_content_paragraph_after_initial_blocks = False
            continue

        # 4. Gestione "Oggetto:" in una tabella
        subject_match = subject_line_pattern.match(original_line) # Usa original_line per catturare spazi iniziali
        if subject_match:
            # Estraiamo l'etichetta e il contenuto. L'LLM dovrebbe fornire "Oggetto: Contenuto..."
            # Potrebbe esserci uno spazio extra prima di "Oggetto:", lo rimuoviamo per l'etichetta.
            full_subject_text = subject_match.group(0).strip() # Es. "Oggetto: Ass.to Campagnolo..."
            label_text = "Oggetto:"
            content_after_oggetto = full_subject_text[len(label_text):].strip()

            table = document.add_table(rows=1, cols=2)
            table.autofit = False
            
            # Calcola larghezze: una più stretta per "Oggetto:", l'altra per il resto
            # Margini standard sono 1 pollice per lato. Pagina A4 è 8.27 pollici.
            page_width_inches = document.sections[0].page_width.inches
            left_margin_inches = document.sections[0].left_margin.inches
            right_margin_inches = document.sections[0].right_margin.inches
            content_area_width = page_width_inches - left_margin_inches - right_margin_inches
            
            col0_width = Inches(1.0) # Larghezza per "Oggetto:" (da PDF sembra circa 1 pollice)
            col1_width = Inches(content_area_width - 1.0 - 0.1) # Rimanente, con un piccolo buffer
            
            table.columns[0].width = col0_width
            table.columns[1].width = col1_width
            
            # remove_table_borders(table) # Rendi i bordi invisibili

            cell_label = table.cell(0, 0)
            para_label = cell_label.paragraphs[0]
            para_label.text = label_text
            para_label.paragraph_format.space_before = Pt(0)
            para_label.paragraph_format.space_after = Pt(0)
            for run in para_label.runs:
                run.font.name = settings.DOCX_FONT_NAME
                run.font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL)
            cell_label.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell_label, end=72) # Aggiunge un piccolo margine a destra di "Oggetto:" (0.05 inches * 1440)

            cell_content = table.cell(0, 1)
            para_content = cell_content.paragraphs[0]
            para_content.text = content_after_oggetto
            para_content.paragraph_format.space_before = Pt(0)
            para_content.paragraph_format.space_after = Pt(0)
            for run in para_content.runs:
                run.font.name = settings.DOCX_FONT_NAME
                run.font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL)
            cell_content.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            
            # Aggiungi uno spazio *dopo* la tabella dell'oggetto
            p_after_table = document.add_paragraph()
            p_after_table.paragraph_format.space_before = Pt(0) # Nessuno spazio prima se quello normale è 6pt
            p_after_table.paragraph_format.space_after = Pt(settings.DOCX_SPACE_AFTER_PARAGRAPH if settings.DOCX_SPACE_AFTER_PARAGRAPH > 0 else 6) # Spazio dopo l'oggetto
            is_first_content_paragraph_after_initial_blocks = True
            continue

        # 5. Gestione Titoli di Sezione (es. "1 – DATI GENERALI")
        section_title_match = section_title_pattern.match(stripped_line)
        if section_title_match:
            p = document.add_paragraph(stripped_line)
            fmt = p.paragraph_format
            # Nessun grassetto, font standard da stile 'Normal'
            # Il PDF esempio ha i titoli leggermente più grandi o stesso font ma grassetto.
            # Se vuoi lo stesso font size, ma grassetto:
            # for run in p.runs: run.bold = True
            # Ma hai detto niente grassetto automatico. Quindi, verranno con font normale.
            # Per replicare il PDF, i titoli sono in un font leggermente diverso o bold.
            # Se l'LLM non li mette in grassetto, qui potresti decidere di farlo:
            for run in p.runs: # Applichiamo font leggermente più grande o bold per i titoli
                # run.font.size = Pt(settings.DOCX_FONT_SIZE_HEADING) # Se hai una dimensione diversa
                run.bold = True # Se decidi che i titoli DEVONO essere bold
            fmt.space_before = Pt(12 if not is_first_content_paragraph_after_initial_blocks else 6) # Più spazio prima dei titoli di sezione
            fmt.space_after = Pt(6) # Spazio dopo il titolo
            is_first_content_paragraph_after_initial_blocks = False
            continue

        # 6. Gestione Paragrafi Normali, Liste e Contenuto Indentato
        if not stripped_line and line.strip() == "": # Riga vuota per spaziatura
            if line_num > 0 and lines[line_num-1].strip():
                 document.add_paragraph() # Crea lo spazio dato da \n\n (usa space_after di default)
            is_first_content_paragraph_after_initial_blocks = True # Potrebbe essere una separazione prima di un nuovo blocco
            continue
        elif stripped_line:
            p = document.add_paragraph()
            fmt = p.paragraph_format
            
            # Logica per indentazione (basata su spazi iniziali forniti da LLM)
            num_leading_spaces = len(original_line) - len(original_line.lstrip(' '))

            if original_line.startswith("Dettaglio spedizione        :"): # Etichetta
                p.text = stripped_line
                fmt.space_before = Pt(3) # Meno spazio prima di "Dettaglio spedizione"
                fmt.space_after = Pt(0) # Nessuno spazio dopo l'etichetta, il contenuto segue
                # L'LLM dovrebbe mettere i due punti e il primo item sulla stessa riga
            elif num_leading_spaces > 20 and not stripped_line.startswith("Dettaglio spedizione"): # Continuazione di Dettaglio Spedizione
                p.text = stripped_line
                fmt.left_indent = Inches(0.8) # Valore da aggiustare per allineare con testo dopo ":"
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
                fmt.line_spacing = 1.0
            elif stripped_line.startswith("- ") or re.match(r"^\d+\.\s+", stripped_line): # Item di lista
                p.text = stripped_line
                fmt.left_indent = Inches(0.5)
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)
            else: # Paragrafo standard
                p.text = stripped_line
                if is_first_content_paragraph_after_initial_blocks and not section_title_match:
                    # Aggiungi spazio prima del primo paragrafo narrativo SE non è un titolo di sezione
                    # E se non è immediatamente dopo l'oggetto (che ha già il suo spazio dopo)
                    prev_line_empty_or_subject = (line_num > 0 and not lines[line_num-1].strip()) or \
                                                 (line_num > 0 and subject_line_pattern.match(lines[line_num-1].strip()))
                    if prev_line_empty_or_subject :
                        fmt.space_before = Pt(6) # Spazio prima del primo paragrafo effettivo
                is_first_content_paragraph_after_initial_blocks = False
            continue

    # --- Footer Statico con Testo e Numero di Pagina in Tabella ---
    footer_static_text_line1 = "BN Surveys Srls – Via T.M. Canepari, 20A/5 – 16159 Genova (Italia) – Tel 010.09080217"
    footer_static_text_line2 = "Codice Fiscale e P IVA 02495290997 – REA GE – 490593 – PEC bnsurveys@legalmail.it"

    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False # Assicura footer specifico per questa sezione

    # Calcola la larghezza dell'area di contenuto per il footer
    page_width_inches = section.page_width.inches
    left_margin_inches = section.left_margin.inches
    right_margin_inches = section.right_margin.inches
    content_area_width_inches = page_width_inches - left_margin_inches - right_margin_inches

    # Aggiungi tabella 1 riga, 2 colonne al footer
    # La prima cella conterrà il testo statico, la seconda il numero di pagina
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(content_area_width_inches))
    footer_table.autofit = False
    # remove_table_borders(footer_table) # Bordi invisibili

    # Calcola larghezze colonne footer
    # Rendiamo la colonna del numero pagina più stretta
    page_num_col_width = Inches(0.75)
    # Usa la content_area_width_inches calcolata per il footer
    text_col_width = Inches(content_area_width_inches - page_num_col_width.inches - 0.1) # -0.1 per buffer

    footer_table.columns[0].width = text_col_width
    footer_table.columns[1].width = page_num_col_width
    
    # Cella per testo statico (sinistra)
    cell_text = footer_table.cell(0, 0)
    set_cell_margins(cell_text, top=30, bottom=30) # Piccoli margini verticali
    
    # Paragrafo 1 del testo statico
    p_text1 = cell_text.paragraphs[0] if cell_text.paragraphs else cell_text.add_paragraph()
    p_text1.clear() # Pulisci eventuale paragrafo di default della cella
    
    parts1 = footer_static_text_line1.split("BN Surveys Srls", 1)
    if len(parts1) == 2:
        if parts1[0]: p_text1.add_run(parts1[0])
        run_company1 = p_text1.add_run("BN Surveys Srls")
        run_company1.font.color.rgb = FOOTER_BN_SURVEYS_COLOR
        p_text1.add_run(parts1[1])
    else:
        p_text1.add_run(footer_static_text_line1)
        
    for run in p_text1.runs:
        run.font.name = settings.DOCX_FONT_NAME
        run.font.size = Pt(12) # Font piccolo per footer
    p_text1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_text1.paragraph_format.space_before = Pt(0)
    p_text1.paragraph_format.space_after = Pt(0)
    p_text1.paragraph_format.line_spacing = 1.0

    # Paragrafo 2 del testo statico (nella stessa cella)
    p_text2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_text2 = cell_text.add_paragraph()
    # Non c'è "BN Surveys Srls" nella seconda riga del footer come da tuo esempio
    p_text2.add_run(footer_static_text_line2)
    for run in p_text2.runs:
        run.font.name = settings.DOCX_FONT_NAME
        run.font.size = Pt(12)
    p_text2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_text2.paragraph_format.space_before = Pt(0)
    p_text2.paragraph_format.space_after = Pt(0)
    p_text2.paragraph_format.line_spacing = 1.0
    
    cell_text.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM # Allinea il testo in basso nella cella

    # Cella per numero di pagina (destra)
    cell_page_num = footer_table.cell(0, 1)
    set_cell_margins(cell_page_num, top=30, bottom=30) # Piccoli margini verticali
    
    p_page_num = cell_page_num.paragraphs[0] if cell_page_num.paragraphs else cell_page_num.add_paragraph()
    p_page_num.clear()
    p_page_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_page_num.paragraph_format.space_before = Pt(0)
    p_page_num.paragraph_format.space_after = Pt(0)
    
    run_page = p_page_num.add_run()
    run_page.font.name = settings.DOCX_FONT_NAME
    run_page.font.size = Pt(12) # Come da PDF
    
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE" 
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    
    run_page._r.append(fldChar_begin)
    run_page._r.append(instrText)
    run_page._r.append(fldChar_end)
    
    cell_page_num.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP # Allinea in alto


    # Salva in un oggetto BytesIO
    file_stream: io.BytesIO = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream