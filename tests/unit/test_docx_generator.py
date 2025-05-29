"""
Unit tests for DOCX generator functionality.
Tests the final step in the user flow: converting plain text reports to styled DOCX files.
"""
import pytest
from unittest.mock import Mock, patch, MagicMock
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import the function under test
from docx_generator import create_styled_docx


class TestDocxGeneration:
    """Test DOCX file generation from plain text reports."""
    
    def test_create_styled_docx_basic_functionality(self):
        """Test basic DOCX creation with simple plain text."""
        # Arrange
        plain_text = "This is a simple test report.\n\nSecond paragraph."
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        assert isinstance(result, io.BytesIO), "Should return BytesIO object"
        assert result.tell() == 0, "Stream should be at beginning"
        
        # Verify it's a valid DOCX by opening it
        document = Document(result)
        assert len(document.paragraphs) > 0, "Document should have paragraphs"
    
    def test_create_styled_docx_with_section_titles(self):
        """Test DOCX creation with section titles that should be bolded."""
        # Arrange
        plain_text = """OGGETTO: Test Report

1. DATI DELLA SPEDIZIONE
This is content under the first section.

2. DINAMICA DEGLI EVENTI ED ACCERTAMENTI
This is content under the second section.

COMPUTO DEL DANNO:
This is damage calculation content."""
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        
        # Find paragraphs that should be section titles
        section_titles = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip() in ["OGGETTO: Test Report", "1. DATI DELLA SPEDIZIONE", 
                                         "2. DINAMICA DEGLI EVENTI ED ACCERTAMENTI", "COMPUTO DEL DANNO:"]:
                section_titles.append(paragraph)
        
        assert len(section_titles) > 0, "Should find section title paragraphs"
        
        # Check that section titles are bold
        for title_paragraph in section_titles:
            if title_paragraph.runs:
                assert title_paragraph.runs[0].bold, f"Section title '{title_paragraph.text}' should be bold"
    
    def test_create_styled_docx_with_list_items(self):
        """Test DOCX creation with list items that should be indented."""
        # Arrange
        plain_text = """Regular paragraph.

- First bullet point
- Second bullet point
* Alternative bullet style
1. Numbered item one
2. Numbered item two

Another regular paragraph."""
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        
        # Find paragraphs that should be list items
        list_items = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if (text.startswith('- ') or text.startswith('* ') or 
                (text and text[0].isdigit() and '. ' in text[:5])):
                list_items.append(paragraph)
        
        assert len(list_items) > 0, "Should find list item paragraphs"
        
        # Check that list items have left indentation
        for list_paragraph in list_items:
            assert list_paragraph.paragraph_format.left_indent is not None, f"List item '{list_paragraph.text}' should be indented"
    
    def test_create_styled_docx_empty_content(self):
        """Test DOCX creation with empty content."""
        # Arrange
        plain_text = ""
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        assert isinstance(result, io.BytesIO), "Should return BytesIO object even for empty content"
        
        document = Document(result)
        # Should still have header and footer elements
        assert len(document.sections) > 0, "Should have at least one section"
    
    def test_create_styled_docx_with_multiple_blank_lines(self):
        """Test DOCX creation handles multiple consecutive blank lines properly."""
        # Arrange
        plain_text = """First paragraph.



Second paragraph after multiple blank lines.


Third paragraph."""
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        
        # Should handle blank lines without creating excessive empty paragraphs
        non_empty_paragraphs = [p for p in document.paragraphs if p.text.strip()]
        assert len(non_empty_paragraphs) >= 3, "Should preserve content paragraphs"


class TestDocxStyling:
    """Test DOCX styling and formatting features."""
    
    @patch('docx_generator.settings')
    def test_font_configuration_applied(self, mock_settings):
        """Test that font settings from configuration are applied."""
        # Arrange
        mock_settings.DOCX_FONT_NAME = 'Arial'
        mock_settings.DOCX_FONT_SIZE_NORMAL = 12
        mock_settings.DOCX_FONT_SIZE_HEADING = 14
        
        plain_text = "Test content"
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        normal_style = document.styles['Normal']
        assert normal_style.font.name == 'Arial', "Should use configured font name"
        assert normal_style.font.size == Pt(12), "Should use configured font size"
    
    @patch('docx_generator.settings')
    def test_header_configuration_applied(self, mock_settings):
        """Test that header settings from configuration are applied."""
        # Arrange
        mock_settings.DOCX_HEADER_TEXT = "TEST COMPANY HEADER"
        mock_settings.DOCX_FONT_SIZE_HEADING = 16
        mock_settings.DOCX_FONT_NAME = 'Times New Roman'
        mock_settings.DOCX_FONT_SIZE_NORMAL = 11
        mock_settings.DOCX_FOOTER_TEXT_TEMPLATE = "Footer - Pagina {page_number} di {total_pages}"
        
        plain_text = "Test content"
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        header = document.sections[0].header
        
        # Check header content
        header_text = ""
        for paragraph in header.paragraphs:
            header_text += paragraph.text
        
        assert "TEST COMPANY HEADER" in header_text, "Should include configured header text"
    
    def test_date_formatting_in_document(self):
        """Test that current date is properly formatted and included."""
        # Arrange
        plain_text = "Test content"
        
        # Act
        with patch('docx_generator.datetime') as mock_datetime:
            mock_now = Mock()
            mock_now.strftime.return_value = "04 luglio 2024"
            mock_datetime.now.return_value = mock_now
            
            result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        
        # Find paragraph with date
        date_found = False
        for paragraph in document.paragraphs:
            if "Genova," in paragraph.text and "luglio" in paragraph.text:
                date_found = True
                assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT, "Date should be right-aligned"
                break
        
        assert date_found, "Should include formatted date in document"


class TestDocxStructure:
    """Test DOCX document structure and static elements."""
    
    @patch('docx_generator.settings')
    def test_static_header_elements_included(self, mock_settings):
        """Test that static header elements are included in the document."""
        # Arrange
        mock_settings.DOCX_HEADER_TEXT = "Salomone & ASSOCIATI SRL"
        mock_settings.DOCX_FONT_SIZE_HEADING = 14
        mock_settings.DOCX_FONT_NAME = 'Arial'
        mock_settings.DOCX_FONT_SIZE_NORMAL = 11
        mock_settings.DOCX_FOOTER_TEXT_TEMPLATE = "Footer text"
        
        plain_text = "Test content"
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        
        # Check for recipient address block placeholder
        recipient_found = False
        for paragraph in document.paragraphs:
            if "Spett.le" in paragraph.text and "Client Name Placeholder" in paragraph.text:
                recipient_found = True
                break
        
        assert recipient_found, "Should include recipient address block placeholder"
        
        # Check for internal references block
        references_found = False
        for paragraph in document.paragraphs:
            if "Vs. rif.:" in paragraph.text and "Rif. broker:" in paragraph.text:
                references_found = True
                break
        
        assert references_found, "Should include internal references block"
    
    @patch('docx_generator.settings')
    def test_footer_with_page_numbering(self, mock_settings):
        """Test that footer includes page numbering fields."""
        # Arrange
        mock_settings.DOCX_FOOTER_TEXT_TEMPLATE = "Pagina {page_number} di {total_pages}"
        mock_settings.DOCX_HEADER_TEXT = "Header"
        mock_settings.DOCX_FONT_NAME = 'Arial'
        mock_settings.DOCX_FONT_SIZE_NORMAL = 11
        mock_settings.DOCX_FONT_SIZE_HEADING = 14
        
        plain_text = "Test content"
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        footer = document.sections[0].footer
        
        assert len(footer.paragraphs) > 0, "Footer should have paragraphs"
        
        # Check footer alignment
        footer_paragraph = footer.paragraphs[0]
        assert footer_paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER, "Footer should be center-aligned"
    
    def test_document_sections_structure(self):
        """Test that document has proper section structure."""
        # Arrange
        plain_text = "Test content"
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        
        assert len(document.sections) >= 1, "Document should have at least one section"
        
        section = document.sections[0]
        assert section.header is not None, "Section should have header"
        assert section.footer is not None, "Section should have footer"


class TestDocxErrorHandling:
    """Test error handling in DOCX generation."""
    
    def test_create_styled_docx_with_special_characters(self):
        """Test DOCX creation with special characters and unicode."""
        # Arrange
        plain_text = """Report with special characters:
        
€ 1.234,56 - Euro amount
Società Assicuratrice
Più informazioni
Qualità del servizio
Città: Milano"""
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        assert isinstance(result, io.BytesIO), "Should handle special characters without error"
        
        document = Document(result)
        
        # Verify special characters are preserved
        content_found = False
        for paragraph in document.paragraphs:
            if "€" in paragraph.text or "Società" in paragraph.text:
                content_found = True
                break
        
        assert content_found, "Should preserve special characters in content"
    
    def test_create_styled_docx_with_very_long_content(self):
        """Test DOCX creation with very long content."""
        # Arrange
        long_paragraph = "This is a very long paragraph. " * 100
        plain_text = f"""OGGETTO: Long Report

1. SECTION ONE
{long_paragraph}

2. SECTION TWO  
{long_paragraph}

CONCLUSION:
{long_paragraph}"""
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        assert isinstance(result, io.BytesIO), "Should handle long content without error"
        
        document = Document(result)
        assert len(document.paragraphs) > 0, "Should create document with long content"
    
    @patch('docx_generator.Document')
    def test_create_styled_docx_document_creation_error(self, mock_document_class):
        """Test error handling when document creation fails."""
        # Arrange
        mock_document_class.side_effect = Exception("Document creation failed")
        plain_text = "Test content"
        
        # Act & Assert
        with pytest.raises(Exception) as exc_info:
            create_styled_docx(plain_text)
        
        assert "Document creation failed" in str(exc_info.value)


class TestRegexPatterns:
    """Test regex patterns used for text formatting."""
    
    def test_section_title_pattern_matching(self):
        """Test that section title regex correctly identifies titles."""
        # Arrange
        plain_text = """OGGETTO: Test Report

1. DATI DELLA SPEDIZIONE
Regular paragraph content.

2. DINAMICA DEGLI EVENTI ED ACCERTAMENTI  
More content here.

COMPUTO DEL DANNO:
Damage calculation.

NOT A TITLE: lowercase content
also not a title"""
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        
        # Check that proper titles are identified and formatted
        bold_paragraphs = []
        for paragraph in document.paragraphs:
            if paragraph.runs and paragraph.runs[0].bold:
                bold_paragraphs.append(paragraph.text.strip())
        
        expected_titles = [
            "OGGETTO: Test Report",
            "1. DATI DELLA SPEDIZIONE", 
            "2. DINAMICA DEGLI EVENTI ED ACCERTAMENTI",
            "COMPUTO DEL DANNO:"
        ]
        
        for title in expected_titles:
            assert any(title in bold_text for bold_text in bold_paragraphs), f"Title '{title}' should be bold"
    
    def test_list_item_pattern_matching(self):
        """Test that list item regex correctly identifies list items."""
        # Arrange
        plain_text = """Regular paragraph.

- Bullet point one
- Bullet point two
* Alternative bullet
1. Numbered item
2. Another numbered item
10. Double digit number

Not a list item
- But this is a list item"""
        
        # Act
        result = create_styled_docx(plain_text)
        
        # Assert
        document = Document(result)
        
        # Check that list items have indentation
        indented_paragraphs = []
        for paragraph in document.paragraphs:
            if (paragraph.paragraph_format.left_indent is not None and 
                paragraph.paragraph_format.left_indent > Inches(0)):
                indented_paragraphs.append(paragraph.text.strip())
        
        expected_list_items = [
            "- Bullet point one",
            "- Bullet point two", 
            "* Alternative bullet",
            "1. Numbered item",
            "2. Another numbered item",
            "10. Double digit number",
            "- But this is a list item"
        ]
        
        for item in expected_list_items:
            assert any(item in indented_text for indented_text in indented_paragraphs), f"List item '{item}' should be indented" 